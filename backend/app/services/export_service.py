"""Export service — PDF/DOCX/MD 导出。"""
from __future__ import annotations
from pathlib import Path
from app.config import EXPORT_DIR


def export_to_markdown(content: str, title: str = "document") -> str:
    """导出 MD,返回文件路径。"""
    path = EXPORT_DIR / f"{title}.md"
    path.write_text(content, encoding="utf-8")
    return str(path)


def export_to_pdf(content: str, title: str = "document") -> str:
    """PDF 导出 (ReportLab 流水式排版, 支持 CJK)。"""
    path = _export_pdf_reportlab(content, title)
    if path and Path(path).stat().st_size > 0:
        with open(path, "rb") as f:
            if f.read(5) == b"%PDF-":
                return path
    from loguru import logger
    logger.warning("[export] reportlab PDF 无效, 回退到 MD")
    return export_to_markdown(content, title)


def _export_pdf_weasyprint(content: str, title: str = "document") -> str | None:
    """WeasyPrint HTML→PDF 导出。"""
    try:
        import markdown
        from weasyprint import HTML as WeasyDoc
    except (ImportError, OSError):
        return None

    try:
        html_body = markdown.markdown(content, extensions=["extra", "tables"])
        styled_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<style>
  @page {{ size: A4; margin: 2cm; }}
  body {{ font-family: -apple-system, "PingFang SC", "ArialUni", sans-serif; font-size: 11pt; line-height: 1.7; }}
  h1 {{ font-size: 18pt; margin: 16pt 0 8pt; }}
  h2 {{ font-size: 15pt; margin: 14pt 0 6pt; }}
  h3 {{ font-size: 13pt; margin: 12pt 0 6pt; }}
  table {{ border-collapse: collapse; width: 100%; margin: 8pt 0; font-size: 9pt; }}
  td, th {{ border: 1px solid #ccc; padding: 4px 6px; }}
  th {{ background: #f0f0f0; }}
</style>
</head>
<body>{html_body}</body>
</html>"""
        path = str(EXPORT_DIR / f"{title}.pdf")
        WeasyDoc(string=styled_html).write_pdf(path)
        return path
    except Exception as e:
        from loguru import logger
        logger.warning(f"[export] weasyprint 失败: {e}")
        return None


def _export_pdf_reportlab(content: str, title: str = "document") -> str:
    """ReportLab 流水式 PDF — 支持 CJK,自动分页。"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import re

        # 注册中文字体
        font_path = "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"
        pdfmetrics.registerFont(TTFont("ArialUni", font_path))

        styles = getSampleStyleSheet()

        style_normal = ParagraphStyle(
            "CNNormal", fontName="ArialUni", fontSize=10, leading=16,
            spaceAfter=4, wordWrap="CJK",
        )
        style_h1 = ParagraphStyle("CNH1", fontName="ArialUni", fontSize=16, leading=24, spaceBefore=12, spaceAfter=6)
        style_h2 = ParagraphStyle("CNH2", fontName="ArialUni", fontSize=14, leading=20, spaceBefore=10, spaceAfter=4)
        style_h3 = ParagraphStyle("CNH3", fontName="ArialUni", fontSize=12, leading=18, spaceBefore=8, spaceAfter=4)
        style_code = ParagraphStyle("CNCode", fontName="ArialUni", fontSize=8, leading=12, spaceAfter=4, leftIndent=10)

        story = []
        lines = content.split("\n")
        i = 0
        while i < len(lines):
            raw = lines[i]
            txt = raw.strip()

            if not txt:
                story.append(Spacer(1, 4))
                i += 1
                continue
            if txt in ("---", "***", "___"):
                story.append(Spacer(1, 6))
                i += 1
                continue

            # 去除 markdown 标记
            clean = txt
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", clean)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
            clean = re.sub(r"`(.+?)`", r"\1", clean)
            # escape XML special chars for Paragraph
            clean = clean.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

            if clean.startswith("#### "):
                story.append(Paragraph(clean[5:], style_normal))
            elif clean.startswith("### "):
                story.append(Paragraph(clean[4:], style_h3))
            elif clean.startswith("## "):
                story.append(Paragraph(clean[3:], style_h2))
            elif clean.startswith("# "):
                story.append(Paragraph(clean[2:], style_h1))
            elif clean.startswith("```"):
                # 代码块: 收集直到结束
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                code_text = "<br/>".join(
                    l.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    for l in code_lines
                )
                if code_text:
                    story.append(Paragraph(code_text, style_code))
            elif clean.startswith("|"):
                # 表格行
                cells = [c.strip() for c in clean.strip("|").split("|")]
                if not any(c.replace("-", "").replace(":", "").strip() for c in cells):
                    # 分隔行, 跳过
                    i += 1
                    continue
                # 收集所有表格行
                table_rows = [cells]
                i += 1
                while i < len(lines) and lines[i].strip().startswith("|"):
                    row_clean = lines[i].strip().strip("|")
                    row_cells = [c.strip() for c in row_clean.split("|")]
                    is_sep = not any(c.replace("-", "").replace(":", "").strip() for c in row_cells)
                    if not is_sep:
                        table_rows.append(row_cells)
                    i += 1
                if len(table_rows) >= 2:
                    t = Table(table_rows, hAlign="LEFT")
                    t.setStyle(TableStyle([
                        ("FONTNAME", (0, 0), (-1, -1), "ArialUni"),
                        ("FONTSIZE", (0, 0), (-1, -1), 8),
                        ("GRID", (0, 0), (-1, -1), 0.5, (0.5, 0.5, 0.5)),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 2),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 6))
                continue
            else:
                story.append(Paragraph(clean, style_normal))

            i += 1

        path = str(EXPORT_DIR / f"{title}.pdf")
        SimpleDocTemplate(
            path, pagesize=A4,
            leftMargin=2*cm, rightMargin=2*cm,
            topMargin=2*cm, bottomMargin=2*cm,
        ).build(story)
        return path
    except Exception as e:
        from loguru import logger
        logger.warning(f"[export] reportlab 失败, 回退到 MD: {e}")
        return export_to_markdown(content, title)


def export_to_docx(content: str, title: str = "document") -> str:
    """简单 DOCX 导出(纯文本,后续可丰富)。"""
    try:
        from docx import Document
        doc = Document()
        for line in content.split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.strip():
                doc.add_paragraph(line)
        path = EXPORT_DIR / f"{title}.docx"
        doc.save(str(path))
        return str(path)
    except Exception:
        # fallback to MD
        return export_to_markdown(content, title)